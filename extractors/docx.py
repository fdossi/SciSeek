"""
Extractor para arquivos DOCX usando python-docx.
"""

import time
from pathlib import Path
from typing import Dict, Any, Tuple
import unicodedata
from docx import Document
from .base import BaseExtractor, ExtractionResult


class DocxExtractor(BaseExtractor):
    """Extrai texto de arquivos DOCX (Word)."""

    def _get_supported_extensions(self) -> Tuple[str, ...]:
        return (".docx", ".doc")

    def can_process(self, file_path: Path) -> bool:
        """Verifica se pode processar o arquivo DOCX."""
        # Para .doc, verifica apenas extensão (não é totalmente suportado)
        return file_path.suffix.lower() in self.supported_extensions

    def extract(self, file_path: Path) -> ExtractionResult:
        """
        Extrai texto de um arquivo DOCX.
        
        Args:
            file_path: Caminho para o arquivo DOCX
            
        Returns:
            ExtractionResult com o conteúdo extraído
        """
        start_time = time.time()

        # Validar arquivo
        is_valid, error_msg = self.validate_file(file_path)
        if not is_valid:
            return ExtractionResult(
                file_path=file_path,
                file_name=file_path.name,
                format_type="docx",
                text_content="",
                metadata={},
                error=error_msg,
                success=False
            )

        try:
            doc = Document(str(file_path))
            
            # Extrair texto dos parágrafos
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extrair texto das tabelas (opcional, conforme config)
            include_tables = self.config.get("include_tables", True)
            if include_tables:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                paragraphs.append(cell.text)

            combined_text = "\n".join(paragraphs)
            normalized_text = self._normalize_text(combined_text)

            # Extrair metadados
            metadata = self._extract_metadata(doc)

            extraction_time_ms = (time.time() - start_time) * 1000

            return ExtractionResult(
                file_path=file_path,
                file_name=file_path.name,
                format_type="docx",
                text_content=normalized_text,
                metadata=metadata,
                extraction_time_ms=extraction_time_ms,
                success=True
            )

        except Exception as e:
            extraction_time_ms = (time.time() - start_time) * 1000
            return ExtractionResult(
                file_path=file_path,
                file_name=file_path.name,
                format_type="docx",
                text_content="",
                metadata={},
                extraction_time_ms=extraction_time_ms,
                error=f"Erro ao extrair DOCX: {str(e)}",
                success=False
            )

    def _normalize_text(self, text: str) -> str:
        """Normaliza o texto extraído."""
        text = unicodedata.normalize("NFKD", text)
        text = " ".join(text.split())
        return text

    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """Extrai metadados do DOCX."""
        metadata = {}

        try:
            if doc.core_properties:
                props = doc.core_properties
                metadata["title"] = props.title or ""
                metadata["author"] = props.author or ""
                metadata["subject"] = props.subject or ""
                metadata["created"] = str(props.created) if props.created else ""
                metadata["modified"] = str(props.modified) if props.modified else ""
        except Exception:
            pass

        return metadata
